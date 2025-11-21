import streamlit as st
import requests
import json
import re
import io
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
from datetime import datetime

st.set_page_config(page_title="GLR Pipeline", layout="wide")

st.title("📋 GLR Pipeline - Insurance Template Automation")
st.markdown("**Automate insurance template filling using photo reports and LLMs**")

# API Configuration
OPENROUTER_API_KEY = st.sidebar.text_input("OpenRouter API Key", type="password", 
                                            help="Get your API key from openrouter.ai")
MODEL = st.sidebar.selectbox("Select LLM Model", [
    "deepseek/deepseek-chat",
    "meta-llama/llama-3.1-8b-instruct:free",
    "google/gemma-2-9b-it:free",
    "mistralai/mistral-7b-instruct:free"
])

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file."""
    text = ""
    try:
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    except Exception as e:
        st.error(f"Error extracting PDF: {e}")
    return text

def extract_template_fields(template_text):
    """Extract placeholder fields from template."""
    patterns = [
        r'\[([A-Z0-9_]+)\]',
        r'\[XM8_([A-Z0-9_]+)\]',
        r'\{([A-Z0-9_]+)\}',
        r'<<([A-Z0-9_]+)>>',
        r'\[([a-zA-Z0-9_]+)\]'
    ]
    fields = set()
    for pattern in patterns:
        matches = re.findall(pattern, template_text)
        fields.update(matches)
    return list(fields)

def call_llm(prompt, api_key, model):
    """Call OpenRouter LLM API."""
    if not api_key:
        st.error("Please enter your OpenRouter API key in the sidebar.")
        return None
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://streamlit.app",
        "X-Title": "GLR Pipeline"
    }
    
    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4000,
        "temperature": 0.3
    }
    
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"LLM API Error: {e}")
        return None

def extract_key_values_from_reports(photo_reports_text, template_fields, api_key, model):
    """Use LLM to extract key-value pairs from photo reports."""
    
    prompt = f"""You are an insurance claims data extraction assistant. 
    
I have the following photo report text from an insurance inspection:

--- PHOTO REPORT TEXT ---
{photo_reports_text[:15000]}
--- END PHOTO REPORT ---

I need to extract values for these template fields:
{json.dumps(template_fields, indent=2)}

Please analyze the photo report and extract relevant information for each field.
Return ONLY a valid JSON object with field names as keys and extracted values as values.

For fields you cannot find data for, use reasonable defaults or leave as empty string.

Common field mappings:
- INSURED_NAME: The policyholder/insured name
- DATE_LOSS: Date of loss/incident
- DATE_INSPECTED: Date inspection was performed
- INSURED_H_STREET, INSURED_P_STREET: Property address
- INSURED_H_CITY, INSURED_P_CITY: City
- INSURED_H_STATE, INSURED_P_STATE: State
- INSURED_H_ZIP, INSURED_P_ZIP: ZIP code
- CLAIM_NUM: Claim number
- POLICY_NUM: Policy number
- TOL_CODE, TOL_DESC: Type of loss code/description
- MORTGAGEE: Mortgage company
- ESTIMATOR_NAME: Inspector/estimator name

Also extract and summarize:
- Roof damage details (type, slopes affected, number of damaged shingles)
- Elevation damage (front, right, rear, left)
- Interior damage
- Other structures damage
- Fence/pool/other damage

Return ONLY the JSON object, no additional text.
"""
    
    response = call_llm(prompt, api_key, model)
    
    if response:
        try:
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                return json.loads(json_match.group())
        except json.JSONDecodeError:
            st.warning("Could not parse LLM response as JSON. Using raw response.")
    
    return {}

def generate_report_narrative(photo_reports_text, api_key, model):
    """Generate narrative sections for the GLR report."""
    
    prompt = f"""You are an insurance claims report writer. Based on the following photo report inspection data, 
generate professional narrative sections for a General Loss Report (GLR).

--- PHOTO REPORT DATA ---
{photo_reports_text[:12000]}
--- END DATA ---

Generate the following sections in a professional insurance report style:

1. DWELLING DESCRIPTION: Describe the property type, construction, roofing materials
2. PROPERTY CONDITION: General condition observations
3. ROOF INSPECTION: Detailed findings for each slope (front, right, rear, left)
4. FRONT ELEVATION: Damage findings
5. RIGHT ELEVATION: Damage findings
6. REAR ELEVATION: Damage findings
7. LEFT ELEVATION: Damage findings
8. INTERIOR: Any interior damage noted
9. OTHER STRUCTURES: Detached garage, shed, fence, pool findings
10. CAUSE AND ORIGIN: Type of loss and cause

Format your response as JSON with section names as keys and narrative text as values.
Return ONLY the JSON object.
"""
    
    response = call_llm(prompt, api_key, model)
    
    if response:
        try:
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass
    
    return {}

def fill_template(template_doc, extracted_data, narratives):
    """Fill the template document with extracted data."""
    
    for para in template_doc.paragraphs:
        for key, value in extracted_data.items():
            placeholders = [f"[{key}]", f"[XM8_{key}]", f"{{{key}}}", f"<<{key}>>"]
            for placeholder in placeholders:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value) if value else "")
    
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in extracted_data.items():
                    placeholders = [f"[{key}]", f"[XM8_{key}]", f"{{{key}}}", f"<<{key}>>"]
                    for placeholder in placeholders:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value) if value else "")
    
    return template_doc

# Main UI
st.sidebar.markdown("---")
st.sidebar.markdown("""
### 📖 Instructions:
1. Enter your OpenRouter API key
2. Upload a .docx template
3. Upload photo report PDFs
4. Click Process to generate filled report
""")

# File uploaders
col1, col2 = st.columns(2)

with col1:
    st.subheader("📄 Upload Template (.docx)")
    template_file = st.file_uploader("Insurance template document", type=['docx'], key="template")
    
    if template_file:
        st.success(f"✅ Template loaded: {template_file.name}")

with col2:
    st.subheader("📸 Upload Photo Reports (.pdf)")
    photo_files = st.file_uploader("Photo report PDFs", type=['pdf'], 
                                   accept_multiple_files=True, key="photos")
    
    if photo_files:
        st.success(f"✅ {len(photo_files)} photo report(s) loaded")

# Process button
st.markdown("---")

if st.button("🚀 Process and Generate Report", type="primary", 
             disabled=not (template_file and photo_files and OPENROUTER_API_KEY)):
    
    with st.spinner("Processing..."):
        # Step 1: Extract text from photo reports
        st.info("📖 Step 1: Extracting text from photo reports...")
        combined_photo_text = ""
        
        for pdf_file in photo_files:
            pdf_file.seek(0)
            text = extract_text_from_pdf(pdf_file)
            combined_photo_text += f"\n--- {pdf_file.name} ---\n{text}\n"
        
        st.success(f"Extracted {len(combined_photo_text)} characters from photo reports")
        
        # Step 2: Load template and extract fields
        st.info("📋 Step 2: Analyzing template fields...")
        template_file.seek(0)
        template_doc = Document(template_file)
        
        template_text = ""
        for para in template_doc.paragraphs:
            template_text += para.text + "\n"
        
        template_fields = extract_template_fields(template_text)
        st.success(f"Found {len(template_fields)} template fields")
        
        with st.expander("View detected template fields"):
            st.write(template_fields)
        
        # Step 3: Extract key-value pairs using LLM
        st.info("🤖 Step 3: Extracting data using LLM...")
        extracted_data = extract_key_values_from_reports(
            combined_photo_text, template_fields, OPENROUTER_API_KEY, MODEL
        )
        
        if extracted_data:
            st.success(f"Extracted {len(extracted_data)} data points")
            with st.expander("View extracted data"):
                st.json(extracted_data)
        
        # Step 4: Generate narratives
        st.info("✍️ Step 4: Generating report narratives...")
        narratives = generate_report_narrative(
            combined_photo_text, OPENROUTER_API_KEY, MODEL
        )
        
        if narratives:
            st.success("Generated report narratives")
            with st.expander("View generated narratives"):
                st.json(narratives)
        
        # Step 5: Fill template
        st.info("📝 Step 5: Filling template...")
        template_file.seek(0)
        output_doc = Document(template_file)
        
        # Add current date
        extracted_data['DATE_CURRENT'] = datetime.now().strftime("%m/%d/%Y")
        extracted_data['XM8_DATE_CURRENT'] = datetime.now().strftime("%m/%d/%Y")
        
        filled_doc = fill_template(output_doc, extracted_data, narratives)
        
        # Save to buffer
        doc_buffer = io.BytesIO()
        filled_doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        st.success("✅ Report generated successfully!")
        
        # Display preview
        st.markdown("---")
        st.subheader("📄 Generated Report Preview")
        
        preview_text = ""
        for para in filled_doc.paragraphs[:30]:
            if para.text.strip():
                preview_text += para.text + "\n\n"
        
        st.text_area("Report Preview (first 30 paragraphs)", preview_text, height=400)
        
        # Download button
        st.download_button(
            label="📥 Download Filled Report (.docx)",
            data=doc_buffer,
            file_name=f"GLR_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Show narratives for manual insertion
        if narratives:
            st.markdown("---")
            st.subheader("📝 Generated Narratives (for manual insertion)")
            
            for section, content in narratives.items():
                with st.expander(f"📌 {section}"):
                    st.write(content)

else:
    if not OPENROUTER_API_KEY:
        st.warning("⚠️ Please enter your OpenRouter API key in the sidebar")
    if not template_file:
        st.info("📄 Please upload a .docx template file")
    if not photo_files:
        st.info("📸 Please upload photo report PDF file(s)")

# Footer
st.markdown("---")
st.markdown("""
### 🎯 Supported Template Formats:
- USAA GLR Template
- Elevate Claims Template  
- Eberl/GuideOne Template
- Custom templates with `[FIELD_NAME]` or `[XM8_FIELD_NAME]` placeholders

### 🔗 Get API Key:
Visit [OpenRouter.ai](https://openrouter.ai) to get a free API key for DeepSeek and other models.
""")
